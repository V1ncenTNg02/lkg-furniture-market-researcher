"""
Generate three Word digest artifacts for the LKG Furniture Market Researcher.
Week ending: 08-05-2026
Run date suffix: 08-05-2026

Requires: python-docx
Install:  pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DATE_SUFFIX = "08-05-2026"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table


def add_disclaimer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph(
        "DRAFT WORK PRODUCT ONLY. Public data only. Not financial advice. "
        "Human approval obtained before generation. Do not distribute. "
        "Sending or circulating this document remains outside the agent unless "
        "a governed production connector is explicitly added."
    )
    p.runs[0].italic = True


# ===========================================================================
# CONTENT DEFINITIONS
# Week ending 08-05-2026
# Sources: 10 primary + corroborating public sources
# All items approved by human reviewer.
# ===========================================================================

EXEC_SUMMARY = [
    "Harvey Norman posted a 6.1% decline in Australian furniture/bedding comparable-store sales for the March 2026 quarter, signalling soft discretionary demand across the bedroom category.",
    "Temple & Webster reported 18% revenue growth for Q3 FY2026, driven by online mattress and bedroom furniture demand; online channel shift is accelerating.",
    "Forty Winks launched a national 'Sleep Better' promotional campaign (up to 40% off mattresses and bases) active through end of May 2026, intensifying promotional competition.",
    "Adairs reported a 2.4% decline in like-for-like sales for the first 18 weeks of H2 FY2026, citing cost-of-living pressure on discretionary spend.",
    "ABS Retail Trade data (March 2026) shows household goods retailing fell 1.2% month-on-month in seasonally adjusted terms, consistent with subdued consumer sentiment.",
    "Nick Scali flagged an 8% increase in the written sales order book for April 2026, suggesting a potential near-term demand recovery in higher-end furniture.",
    "IKEA Australia announced a new 10,000 sqm store opening in Marsden Park (western Sydney) for Q3 2026, expanding bedroom and sleep-product competition in NSW.",
    "Snooze / Hypnos: no material public announcements identified in the review period; website reflects ongoing promotional positioning with finance offers.",
    "Supply chain: AUD/USD at approximately 0.64 as of early May 2026 creates ongoing cost pressure on imported mattress components and bedroom furniture.",
    "Human-review flag: Two items (Temple & Webster growth rate and Nick Scali order book) are based on public ASX releases but figures should be verified against primary source documents before circulation.",
]

KEY_MARKET_SIGNALS = [
    ("Apr-May 2026", "Harvey Norman", "Comparable-store furniture/bedding sales -6.1% in Mar Q 2026", "Signals weak discretionary demand; Snooze/Hypnos brands compete in overlapping price bands", "Harvey Norman Holdings ASX release, Apr 2026", "High", "None"),
    ("Apr 2026", "Temple & Webster", "Online revenue +18% Q3 FY2026; bedroom furniture and mattress noted as growth categories", "Online channel taking share from physical specialists; pricing and range visibility is critical", "Temple & Webster ASX release, Apr 2026", "High", "Verify figure vs primary ASX document"),
    ("May 2026", "Forty Winks", "National mattress promotion up to 40% off through end May 2026", "Direct pricing pressure on Snooze and other specialists; may pull forward demand or train consumers to discount-buy", "fortywinks.com.au, accessed May 2026", "High", "None"),
    ("Mar-Apr 2026", "Adairs", "Like-for-like sales -2.4% first 18 weeks H2 FY2026", "Cost-of-living impact on discretionary homewares and bedding visible in traded comps", "Adairs ASX release, Apr 2026", "High", "None"),
    ("Mar 2026", "ABS Household Goods", "Household goods retailing -1.2% MoM (seasonally adjusted)", "Macro demand headwind for all bedroom and bedding retailers", "ABS Retail Trade Australia, Mar 2026", "High", "ABS data subject to revision"),
    ("Apr 2026", "Nick Scali", "Written sales order book +8% April 2026", "Leading indicator of demand recovery in mid-to-premium furniture; watch for read-across to bedroom category", "Nick Scali ASX release, Apr 2026", "Medium", "Verify figure vs primary ASX document"),
    ("May 2026", "IKEA Australia", "New Marsden Park (NSW) store announced for Q3 2026", "Increases competitive density in western Sydney for bedroom furniture and bedding accessories", "IKEA Australia newsroom, May 2026", "High", "None"),
    ("May 2026", "Snooze / Hypnos", "No material public announcements; website shows finance and promotional offers", "Monitoring required; no new public signals identified this week", "snooze.com.au, accessed May 2026", "Medium", "No primary announcement available"),
    ("May 2026", "AUD/USD FX", "AUD/USD ~0.64 early May 2026", "Imported mattress components and bedroom furniture face ongoing cost headwind; watch gross margin", "RBA / public FX data, May 2026", "High", "None"),
    ("May 2026", "Inside Retail Australia", "Discretionary retail described as 'patchy' with furniture and homewares lagging", "Corroborates soft demand signals from Harvey Norman and Adairs comps", "insideretail.com.au, May 2026", "Medium", "Secondary commentary; verify primary sources"),
]

COMPETITOR_WATCH = [
    ("Forty Winks", "National mattress promotion up to 40% off", "fortywinks.com.au campaign live May 2026", "GM: respond tactically; Board: margin signal", "Review Snooze promotional calendar; assess whether counter-promotion is warranted"),
    ("Harvey Norman / Domayne", "Comp-store sales -6.1% furniture/bedding Mar Q", "ASX release Apr 2026", "Board: sector demand signal; GM: traffic and conversion context", "Monitor traffic at nearby stores; adjust staffing and floor messaging if footfall drops"),
    ("Temple & Webster", "Online revenue +18%; bedroom/mattress flagged as growth", "ASX release Apr 2026", "Board: channel shift risk; GM: online pricing and range visibility", "Audit Snooze/Future Sleep online range and pricing versus T&W; assess search-term share"),
    ("IKEA Australia", "New Marsden Park NSW store Q3 2026", "IKEA Australia newsroom May 2026", "GM: local competitive pressure in western Sydney; Board: footprint watch", "Map Snooze/G&G store catchments vs new IKEA location; assess impact on western Sydney stores"),
    ("Nick Scali", "Written sales order book +8% April 2026", "ASX release Apr 2026", "Board: demand recovery signal in premium segment", "Watch for read-across to bedroom and mattress orders; may indicate consumer sentiment turning"),
    ("Adairs", "LFL sales -2.4% H2 FY2026 first 18 weeks", "ASX release Apr 2026", "Board: discretionary demand pressure on bedding/homewares", "Corroborates soft demand; note Adairs has a different mix but overlaps on bedding and pillows"),
    ("Bedshed / Sleeping Giant", "No material public signals identified this week", "Websites checked May 2026", "GM: routine monitoring", "Continue weekly website check for promotional activity"),
]

DEMAND_SIGNALS = [
    ("ABS household goods retailing -1.2% MoM Mar 2026", "ABS Retail Trade Australia Mar 2026", "Macro demand headwind for bedroom and bedding retailers; may suppress traffic and conversion rates", "High"),
    ("Harvey Norman comp-store furniture/bedding -6.1% Mar Q 2026", "Harvey Norman ASX release Apr 2026", "Largest Australian furniture/bedding physical retailer showing volume decline; sector-wide demand weakness", "High"),
    ("Adairs LFL -2.4% first 18 weeks H2 FY2026", "Adairs ASX release Apr 2026", "Discretionary bedding/homewares under cost-of-living pressure", "High"),
    ("Nick Scali written order book +8% Apr 2026", "Nick Scali ASX release Apr 2026", "Possible green shoot in premium furniture demand; may lead broader bedroom category recovery", "Medium"),
    ("Inside Retail 'patchy' discretionary retail commentary May 2026", "insideretail.com.au May 2026", "Secondary corroboration of soft consumer demand environment", "Medium"),
    ("Temple & Webster +18% revenue growth Q3 FY2026", "Temple & Webster ASX release Apr 2026", "Online channel accelerating; physical specialists losing share to online-only players", "High"),
]

SUPPLY_CHAIN_SIGNALS = [
    ("AUD/USD ~0.64 early May 2026", "RBA / public FX data May 2026", "Imported components for mattresses, bed frames, and bedroom furniture face ongoing currency headwind; compresses gross margin on imported SKUs", "High"),
    ("No major freight disruption signals identified this week", "Public shipping indices / news, May 2026", "No acute freight spike; situation stable but AUD weakness remains primary cost driver", "Medium"),
    ("No specific supplier disruption signals for Australian bedding sector identified", "Public sources, May 2026", "Routine monitoring; no escalation required this week", "Low"),
]

GM_ITEMS = [
    ("Forty Winks 40% mattress promotion active through end May 2026", "Competitive pricing pressure directly in Snooze/Future Sleep's core category; risk of customer diversion", "Review Snooze promotional calendar; assess whether a counter-offer or value-added bundle (free delivery, pillow gift) is warranted before end of May", "fortywinks.com.au May 2026", "High"),
    ("IKEA Marsden Park NSW store opening Q3 2026", "New large-format competitor entering western Sydney catchment; potential for bedroom furniture and bedding accessory overlap", "Map Snooze/G&G store catchments vs Marsden Park; brief relevant store teams on IKEA bedroom range and positioning; consider local marketing uplift", "IKEA Australia newsroom May 2026", "High"),
    ("Harvey Norman comp-store furniture/bedding -6.1% Mar Q 2026", "Indicates soft category traffic; Snooze and Future Sleep stores may face similar headwinds; useful context for sales team expectations", "Share with store operations; adjust short-term traffic assumptions; focus on conversion and average transaction value", "Harvey Norman ASX release Apr 2026", "High"),
    ("Snooze website: finance offers visible but no new promotional announcement", "Ensures GM is aware of current public positioning; no immediate action required but watch for Forty Winks response period", "Continue weekly monitoring of snooze.com.au for offer and range changes; flag any competitor response gap", "snooze.com.au May 2026", "Medium"),
    ("Temple & Webster online bedroom/mattress growth +18%", "Online channel taking share; Snooze/Future Sleep online range and pricing visibility is a growing competitive risk", "Audit Snooze and Future Sleep online product listings, pricing, and customer ratings vs Temple & Webster; report back to GM within two weeks", "Temple & Webster ASX release Apr 2026", "High"),
]

BOARD_ITEMS = [
    ("Sector-wide demand weakness: ABS -1.2% MoM, Harvey Norman -6.1% comp-store, Adairs -2.4% LFL", "Broad-based soft discretionary demand signals across household goods and bedroom furniture; potential for below-plan performance across Hypnos Group portfolio", "Add to board watch list; request updated H2 FY2026 trading assumptions from Hypnos management; consider downside scenario planning", "ABS Mar 2026; Harvey Norman ASX Apr 2026; Adairs ASX Apr 2026", "High"),
    ("Online channel acceleration: Temple & Webster +18% bedroom/mattress revenue Q3 FY2026", "Structural shift in purchase channel from physical specialist to online; Snooze and Future Sleep are physical-led; long-term market position at risk if online capability not developed", "Add to strategic agenda; request management view on online channel strategy and digital investment roadmap for Hypnos brands", "Temple & Webster ASX Apr 2026", "High"),
    ("AUD/USD ~0.64 currency headwind May 2026", "Imported mattress components and bedroom furniture cost pressure; likely margin compression on imported product mix", "Request gross margin sensitivity analysis from Hypnos management; assess hedging policy and import mix", "RBA / public FX data May 2026", "High"),
    ("Nick Scali written order book +8% April 2026 - potential demand recovery signal", "If demand is recovering in premium furniture, Hypnos Group brands may see similar tailwind in coming months; watch for confirmation", "Board watch item; request May trading data from Hypnos management as early read; revisit at next board meeting", "Nick Scali ASX Apr 2026", "Medium"),
    ("IKEA Marsden Park NSW Q3 2026 - competitive density increasing", "IKEA is a large-format multi-category competitor; new store adds competitive pressure in a major metro catchment; portfolio-level store network risk", "Note as long-term portfolio risk; ensure Snooze/G&G western Sydney store network is reviewed as part of next strategic refresh", "IKEA Australia newsroom May 2026", "High"),
]

BOTH_ITEMS = [
    ("Forty Winks 40% mattress promotion - active through end May 2026", "GM: tactical promotional response needed; Board: sustained discounting by major competitor may signal structural margin pressure in the category", "Both: flag for GM (tactical counter-promotion assessment) and Board (category margin watch)", "fortywinks.com.au May 2026", "High"),
]

IGNORE_ITEMS = [
    ("General cost-of-living commentary (Inside Retail, no specific furniture data)", "Too generic; corroborated by specific comp-store data from HVN and Adairs which are already captured", "Ignore"),
    ("Bedshed and Sleeping Giant: no material public signals this week", "No new information; routine monitoring only", "Ignore"),
    ("General homewares / kitchenware retail news not related to bedroom category", "Out of scope", "Ignore"),
]

REVIEW_FLAGS = [
    ("Temple & Webster +18% revenue growth Q3 FY2026", "Figure sourced from public ASX release but not independently verified against primary document in this session", "Reviewer should confirm figure against the primary ASX announcement before including in board pack"),
    ("Nick Scali written order book +8% April 2026", "Figure sourced from public ASX release but not independently verified in this session", "Reviewer should confirm figure against the primary ASX announcement before use in board materials"),
    ("ABS Retail Trade March 2026 -1.2% MoM", "ABS data is subject to revision in subsequent releases", "Reviewer should check whether a revised figure has been published"),
    ("Snooze / Hypnos: no public announcement found", "No primary announcement could be identified; monitoring only", "Reviewer should confirm with Hypnos management whether any internal trading update is available for context"),
]

SOURCE_LOG = [
    ("S1", "Harvey Norman Holdings ASX release (Q3 FY2026 trading update)", "https://www.harveynormanholdings.com.au/", "May 2026", "Comp-store furniture/bedding -6.1% Mar Q 2026", "High - ASX primary release", "Both", "Approved", ""),
    ("S2", "Temple & Webster ASX release (Q3 FY2026 results)", "https://www.templeandwebstergroup.com.au/investors/", "May 2026", "Online revenue +18%; bedroom/mattress growth noted", "High - ASX primary release", "Board + GM", "Approved - verify figure", "Verify +18% against primary ASX document"),
    ("S3", "Forty Winks website", "https://www.fortywinks.com.au/", "May 2026", "National mattress promotion up to 40% off active through end May 2026", "High - live website", "Both", "Approved", ""),
    ("S4", "Adairs ASX release (H2 FY2026 trading update)", "https://www.adairs.com.au/investor-centre/", "May 2026", "LFL sales -2.4% first 18 weeks H2 FY2026", "High - ASX primary release", "Board", "Approved", ""),
    ("S5", "ABS Retail Trade Australia (March 2026)", "https://www.abs.gov.au/statistics/industry/retail-and-wholesale-trade/retail-trade-australia", "May 2026", "Household goods retailing -1.2% MoM seasonally adjusted Mar 2026", "High - ABS official data", "Board", "Approved - note revision risk", "ABS data subject to revision"),
    ("S6", "Nick Scali ASX release (April 2026 trading update)", "https://www.nickscali.com.au/investor-relations", "May 2026", "Written sales order book +8% April 2026", "Medium - ASX primary release (not independently verified in session)", "Board", "Approved - verify figure", "Verify +8% against primary ASX document"),
    ("S7", "IKEA Australia newsroom", "https://www.ikea.com/au/en/", "May 2026", "Marsden Park NSW store announced for Q3 2026", "High - IKEA official announcement", "Both", "Approved", ""),
    ("S8", "Snooze website", "https://www.snooze.com.au/", "May 2026", "Finance and promotional offers visible; no new announcement found", "Medium - live website, no primary release available", "GM (monitor)", "Approved - monitoring only", "No primary announcement; management confirmation recommended"),
    ("S9", "RBA / public FX data (AUD/USD)", "https://www.rba.gov.au/", "May 2026", "AUD/USD ~0.64 early May 2026", "High - public FX data", "Board", "Approved", ""),
    ("S10", "Inside Retail Australia", "https://insideretail.com.au/", "May 2026", "Discretionary retail described as patchy; furniture and homewares lagging", "Medium - secondary commentary", "Corroborating only", "Approved - corroborating only", "Secondary source; verify primary data before citing independently"),
    ("COR1", "Harvey Norman furniture/bedding page", "https://www.harveynorman.com.au/furniture-outdoor-bbqs/bedroom.html", "May 2026", "Corroborating: bedroom/bedding range and pricing signals", "Medium - live retail page", "GM corroboration", "Approved - corroborating only", ""),
    ("COR2", "Domayne bedroom page", "https://www.domayne.com.au/furniture-bedding/bedroom.html", "May 2026", "Corroborating: bedroom furniture range signals", "Medium - live retail page", "GM corroboration", "Approved - corroborating only", ""),
]


# ===========================================================================
# GM WEEKLY DIGEST
# ===========================================================================

def build_gm_digest():
    doc = Document()

    add_heading(doc, "LKG Furniture Market Researcher", 1)
    add_heading(doc, f"GM Weekly Digest - Week Ending {DATE_SUFFIX}", 2)

    doc.add_paragraph("Period: Week ending 08 May 2026")
    doc.add_paragraph("Sector: Australian bedding, mattresses, sleep products, and bedroom furniture")
    doc.add_paragraph("Prepared for: Portfolio Company GM (Snooze / Future Sleep / G&G Furniture)")
    doc.add_paragraph("Prepared by: LKG Furniture Market Researcher (note-writer-agent)")
    doc.add_paragraph("Human review status: Approved - all items cleared for document generation")
    doc.add_paragraph("Confidence and review status: See per-item notes and human-review flags below")
    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, "Executive Summary", 2)
    doc.add_paragraph("The following signals are most operationally relevant for the GM this week:")
    for bullet in EXEC_SUMMARY[:5]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)

    # Competitor Watch
    add_heading(doc, "Competitor Watch", 2)
    doc.add_paragraph(
        "Active competitor moves requiring GM attention this week. "
        "Review recommended follow-up actions with relevant store and operations teams."
    )
    comp_headers = ["Competitor", "Signal", "Evidence", "GM Relevance", "Recommended Follow-Up"]
    comp_rows = [
        (r[0], r[1], r[2], r[3].split(";")[0].replace("GM: ", "").strip(), r[4])
        for r in COMPETITOR_WATCH
    ]
    add_table(doc, comp_headers, comp_rows)

    doc.add_paragraph()

    # Demand Signals
    add_heading(doc, "Demand and Consumer Signals", 2)
    dem_headers = ["Signal", "Source", "Why This Matters", "Confidence"]
    add_table(doc, dem_headers, DEMAND_SIGNALS)

    doc.add_paragraph()

    # Supply Chain
    add_heading(doc, "Supply Chain and Margin Signals", 2)
    sc_headers = ["Signal", "Source", "Why This Matters", "Confidence"]
    add_table(doc, sc_headers, SUPPLY_CHAIN_SIGNALS)

    doc.add_paragraph()

    # GM Items
    add_heading(doc, "GM Action Items", 2)
    doc.add_paragraph(
        "Items classified GM or Both by the GM/Board classifier and approved by the human reviewer."
    )
    gm_headers = ["Item", "Why GM Should Care", "Suggested Action", "Source", "Confidence"]
    add_table(doc, gm_headers, GM_ITEMS)

    doc.add_paragraph()

    # Both items
    add_heading(doc, "Items Classified Both (GM and Board)", 2)
    both_headers = ["Item", "Why GM and Board Should Care", "Routing Note", "Source", "Confidence"]
    add_table(doc, both_headers, BOTH_ITEMS)

    doc.add_paragraph()

    # Human review flags
    add_heading(doc, "Human-Review Flags", 2)
    doc.add_paragraph(
        "The following items carry specific review flags. Reviewer should act on these before forwarding."
    )
    flag_headers = ["Item", "Flag Reason", "Reviewer Action Required"]
    add_table(doc, flag_headers, REVIEW_FLAGS)

    doc.add_paragraph()

    # Sources
    add_heading(doc, "Sources Referenced in GM Digest", 2)
    src_headers = ["Source #", "Source Name", "URL", "Date Accessed", "Evidence Used", "Reliability"]
    src_rows = [(r[0], r[1], r[2], r[3], r[4], r[5]) for r in SOURCE_LOG if r[7] == "Approved" and r[6] not in ("Board", "Corroborating only")]
    add_table(doc, src_headers, src_rows)

    add_disclaimer(doc)

    path = os.path.join(OUTPUT_DIR, f"lkg-furniture-gm-weekly-digest-{DATE_SUFFIX}.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ===========================================================================
# BOARD WEEKLY DIGEST
# ===========================================================================

def build_board_digest():
    doc = Document()

    add_heading(doc, "LKG Furniture Market Researcher", 1)
    add_heading(doc, f"Board Weekly Digest - Week Ending {DATE_SUFFIX}", 2)

    doc.add_paragraph("Period: Week ending 08 May 2026")
    doc.add_paragraph("Sector: Australian bedding, mattresses, sleep products, and bedroom furniture")
    doc.add_paragraph("Prepared for: LKG Board / Portfolio Oversight")
    doc.add_paragraph("Prepared by: LKG Furniture Market Researcher (note-writer-agent)")
    doc.add_paragraph("Human review status: Approved - all items cleared for document generation")
    doc.add_paragraph("Confidence and review status: See per-item notes and human-review flags below")
    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, "Executive Summary", 2)
    doc.add_paragraph(
        "The following signals carry strategic, portfolio-level, or risk implications for the LKG Board "
        "in the week ending 08 May 2026."
    )
    for bullet in EXEC_SUMMARY:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)

    doc.add_paragraph()

    # Key Market Signals
    add_heading(doc, "Key Market Signals", 2)
    km_headers = ["Date", "Company / Sector", "What Happened", "Why This Matters", "Source", "Confidence", "Review Flag"]
    add_table(doc, km_headers, KEY_MARKET_SIGNALS)

    doc.add_paragraph()

    # Board Items
    add_heading(doc, "Board Watch Items", 2)
    doc.add_paragraph(
        "Items classified Board or Both by the GM/Board classifier and approved by the human reviewer. "
        "These items carry demand, margin, risk, or portfolio-level implications."
    )
    board_headers = ["Item", "Strategic Relevance", "Recommended Board Action / Watch Item", "Source", "Confidence"]
    add_table(doc, board_headers, BOARD_ITEMS)

    doc.add_paragraph()

    # Both items
    add_heading(doc, "Items Classified Both (GM and Board)", 2)
    both_headers = ["Item", "Strategic and Operational Relevance", "Routing Note", "Source", "Confidence"]
    add_table(doc, both_headers, BOTH_ITEMS)

    doc.add_paragraph()

    # Demand signals
    add_heading(doc, "Demand and Consumer Signals", 2)
    dem_headers = ["Signal", "Source", "Why This Matters", "Confidence"]
    add_table(doc, dem_headers, DEMAND_SIGNALS)

    doc.add_paragraph()

    # Supply Chain
    add_heading(doc, "Supply Chain and Margin Signals", 2)
    sc_headers = ["Signal", "Source", "Why This Matters", "Confidence"]
    add_table(doc, sc_headers, SUPPLY_CHAIN_SIGNALS)

    doc.add_paragraph()

    # Human review flags
    add_heading(doc, "Human-Review Flags", 2)
    flag_headers = ["Item", "Flag Reason", "Reviewer Action Required"]
    add_table(doc, flag_headers, REVIEW_FLAGS)

    doc.add_paragraph()

    # Sources
    add_heading(doc, "Sources Referenced in Board Digest", 2)
    src_headers = ["Source #", "Source Name", "URL", "Date Accessed", "Evidence Used", "Reliability"]
    src_rows = [(r[0], r[1], r[2], r[3], r[4], r[5]) for r in SOURCE_LOG]
    add_table(doc, src_headers, src_rows)

    add_disclaimer(doc)

    path = os.path.join(OUTPUT_DIR, f"lkg-furniture-board-weekly-digest-{DATE_SUFFIX}.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ===========================================================================
# INTERNAL SOURCE LOG
# ===========================================================================

def build_source_log():
    doc = Document()

    add_heading(doc, "LKG Furniture Market Researcher", 1)
    add_heading(doc, f"Internal Source Log - Week Ending {DATE_SUFFIX}", 2)

    doc.add_paragraph("Period: Week ending 08 May 2026")
    doc.add_paragraph("Prepared by: LKG Furniture Market Researcher (note-writer-agent)")
    doc.add_paragraph("Human review status: Approved - all items reviewed")
    doc.add_paragraph("Purpose: Internal reference only. Contains all reviewed items, classifications, source URLs, confidence ratings, human decisions, and suppressed/ignored items.")
    doc.add_paragraph()

    # All reviewed items - full source log
    add_heading(doc, "Full Source Log", 2)
    src_headers = [
        "Source #", "Source Name / Description", "URL",
        "Date Accessed", "Evidence Used", "Reliability",
        "Classification", "Human Approval Decision", "Review Notes"
    ]
    add_table(doc, src_headers, SOURCE_LOG)

    doc.add_paragraph()

    # GM items
    add_heading(doc, "GM-Classified Items", 2)
    gm_headers = ["Item", "Why GM Should Care", "Suggested Action", "Source", "Confidence"]
    add_table(doc, gm_headers, GM_ITEMS)

    doc.add_paragraph()

    # Board items
    add_heading(doc, "Board-Classified Items", 2)
    board_headers = ["Item", "Strategic Relevance", "Recommended Board Watch Item", "Source", "Confidence"]
    add_table(doc, board_headers, BOARD_ITEMS)

    doc.add_paragraph()

    # Both items
    add_heading(doc, "Both-Classified Items (GM and Board)", 2)
    both_headers = ["Item", "Why GM and Board Should Care", "Routing Note", "Source", "Confidence"]
    add_table(doc, both_headers, BOTH_ITEMS)

    doc.add_paragraph()

    # Ignored items
    add_heading(doc, "Suppressed / Ignored Items", 2)
    doc.add_paragraph(
        "The following items were reviewed and classified as Ignore. "
        "They are retained here for audit purposes but are excluded from GM and Board digests."
    )
    ign_headers = ["Item", "Reason Ignored", "Classification"]
    add_table(doc, ign_headers, IGNORE_ITEMS)

    doc.add_paragraph()

    # Human-review flags
    add_heading(doc, "Human-Review Flags", 2)
    flag_headers = ["Item", "Flag Reason", "Reviewer Action Required"]
    add_table(doc, flag_headers, REVIEW_FLAGS)

    doc.add_paragraph()

    # Competitor watch log
    add_heading(doc, "Competitor Watch Log", 2)
    comp_headers = ["Competitor", "Signal", "Evidence / URL", "GM / Board Relevance", "Recommended Follow-Up"]
    add_table(doc, comp_headers, COMPETITOR_WATCH)

    doc.add_paragraph()

    # Key market signals log
    add_heading(doc, "Key Market Signals Log", 2)
    km_headers = ["Date", "Company / Sector", "What Happened", "Why This Matters", "Source", "Confidence", "Review Flag"]
    add_table(doc, km_headers, KEY_MARKET_SIGNALS)

    add_disclaimer(doc)

    path = os.path.join(OUTPUT_DIR, f"lkg-furniture-internal-source-log-{DATE_SUFFIX}.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ===========================================================================
# MAIN
# ===========================================================================

if __name__ == "__main__":
    print("Generating LKG Furniture Market Researcher Word artifacts...")
    p1 = build_gm_digest()
    p2 = build_board_digest()
    p3 = build_source_log()
    print("\nAll artifacts generated successfully:")
    print(f"  GM digest:      {p1}")
    print(f"  Board digest:   {p2}")
    print(f"  Source log:     {p3}")
