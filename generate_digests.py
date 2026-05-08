"""
LKG Furniture Weekly Digest Generator
Week: May 1-8, 2026
Generates three .docx files for GM, Board, and Internal Source Log audiences.
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r"E:\gitproject\lkg-furniture-market-researcher\output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

WEEK = "May 1-8, 2026"
PREPARED_DATE = "8 May 2026"
DRAFT_DISCLAIMER = (
    "DRAFT - APPROVED FOR INTERNAL DISTRIBUTION ONLY. "
    "Do not forward, publish, or circulate outside Hypnos Group (LK Group) without authorisation. "
    "Prepared: " + PREPARED_DATE + ". Week: " + WEEK + "."
)

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
DARK_NAVY   = RGBColor(0x1F, 0x36, 0x64)   # headings
MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)   # sub-headings
AMBER       = RGBColor(0xC5, 0x50, 0x00)   # review flags / warnings
DARK_GREY   = RGBColor(0x40, 0x40, 0x40)   # body text
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xDD, 0xEA, 0xF1)   # table header fill
ACTION_RED  = RGBColor(0xC0, 0x00, 0x00)   # "Action Required" label


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_heading_style(para, text: str, level: int, doc):
    """Apply colour and bold to a heading paragraph."""
    run = para.runs[0] if para.runs else para.add_run(text)
    if level == 1:
        run.font.size  = Pt(16)
        run.font.color.rgb = DARK_NAVY
        run.bold = True
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = MID_BLUE
        run.bold = True
    elif level == 3:
        run.font.size  = Pt(11)
        run.font.color.rgb = DARK_NAVY
        run.bold = True


def styled_heading(doc, text: str, level: int = 1):
    """Add a styled heading and return the paragraph."""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    style = style_map.get(level, "Heading 1")
    para = doc.add_heading(text, level=level)
    set_heading_style(para, text, level, doc)
    return para


def body_para(doc, text: str = "", bold: bool = False, italic: bool = False,
              colour: RGBColor = None, size: int = 10, space_after: int = 6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = colour or DARK_GREY
    return para


def bullet_para(doc, text: str, bold_prefix: str = None, flag: bool = False,
                indent_level: int = 0):
    """Add a bullet paragraph, optionally with a bold prefix and flag marker."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    if indent_level:
        para.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
    if bold_prefix:
        r = para.add_run(bold_prefix + ": ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_NAVY
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GREY
    if flag:
        r3 = para.add_run("  [REVIEW FLAG]")
        r3.bold = True
        r3.font.color.rgb = AMBER
        r3.font.size = Pt(9)
    return para


def action_box(doc, label: str, text: str):
    """Add an 'Action Required' callout paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Inches(0.2)
    r1 = para.add_run(label + ": ")
    r1.bold = True
    r1.font.color.rgb = ACTION_RED
    r1.font.size = Pt(10)
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GREY
    return para


def add_table_header_row(table, headers: list, bg_hex: str = "1F3664"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        set_cell_bg(cell, bg_hex)
    return row


def disclaimer_footer(doc):
    """Add disclaimer as a small italic paragraph at end."""
    add_horizontal_rule(doc)
    para = doc.add_paragraph()
    run = para.add_run(DRAFT_DISCLAIMER)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    para.paragraph_format.space_before = Pt(8)


def doc_cover(doc, title: str, subtitle: str, audience: str):
    """Add cover block at top of document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK_NAVY

    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)
    r2.font.color.rgb = MID_BLUE

    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Audience: {audience}  |  Week: {WEEK}  |  Prepared: {PREPARED_DATE}")
    r3.font.size = Pt(9)
    r3.italic = True
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_horizontal_rule(doc)
    doc.add_paragraph()


# ===========================================================================
# DOCUMENT 1 — GM WEEKLY DIGEST
# ===========================================================================

def build_gm_digest():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc_cover(
        doc,
        "LKG Furniture — GM Weekly Digest",
        "Hypnos Group (Snooze | Future Sleep | G&G Furniture)",
        "General Manager and Senior Operations Leadership"
    )

    # ------------------------------------------------------------------
    # SECTION 1: Executive Summary
    # ------------------------------------------------------------------
    styled_heading(doc, "1. Executive Summary", level=1)
    bullets = [
        ("Rate headwind confirmed", "RBA's third consecutive hike to 4.35% (5 May) keeps consumer confidence near historic lows (ANZ-Roy Morgan 67.2). Household goods purchase intent at only 15%. Assume caution-mode consumer through Q3 2026."),
        ("A.H. Beard in administration — act now", "Australia's largest domestic mattress manufacturer entered voluntary administration late April. Retail supply chain disruption is live. Competitor stocking gaps and opportunistic range acquisitions are possible in the short window before a buyer is named."),
        ("ACCC $15M Emma Sleep ruling — compliance audit mandatory", "Federal Court ruled Emma Sleep's countdown timers and strikethrough pricing were a 'deliberate marketing strategy.' Snooze must audit all active digital promotional mechanics immediately. Non-compliance risk is real and enforcement is active."),
        ("Two-front promotional battle active now", "Forty Winks 'Forty Frenzy' (up to 50% off mattresses) and Harvey Norman May catalogue are running simultaneously. DTC brands (Koala, Sleeping Duck, Ecosa) are pre-staging EOFY discounts. Margin discipline is the critical lever."),
        ("Margin floor check before 14 May — non-negotiable", "Foam costs up ~20% and ocean freight surcharges $400-900/40ft. Adairs lost 170bps gross margin from equivalent promotional posture in H1 2026. Any promotion launched without an updated COGS check risks trading at negative margin."),
        ("Promotional gap opens 15-31 May", "Competitors wind down between 14-25 May. A service-led, non-discount activation (CRM-based, in-store consultation) can capture motivated buyers without margin compression before EOFY begins."),
        ("PE capital restructuring competitive landscape", "Allegro Funds acquired Fantastic Furniture (86 stores, AUD $563.5M revenue) on 4 May. Amart+Freedom merger (Aug 2025) created a ~120-store entity. Both PE-backed with active investment horizons. Store overlap analysis is required."),
    ]
    for bp, bt in bullets:
        bullet_para(doc, bt, bold_prefix=bp)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 2: Macro & Market Context
    # ------------------------------------------------------------------
    styled_heading(doc, "2. Macro and Market Context", level=1)

    styled_heading(doc, "Interest Rates", level=2)
    body_para(doc, "RBA raised the cash rate to 4.35% on 5 May 2026 — the third consecutive increase. CBA economists expect rates to hold for the remainder of 2026 but leave the door open to further hikes. This is a direct headwind to housing-linked discretionary spending and to consumer willingness to finance large purchases.")
    bullet_para(doc, "Source: RBA Media Release mr-26-12, 5 May 2026", bold_prefix="Source")
    bullet_para(doc, "Review any in-store finance or BNPL offer messaging to ensure it addresses the cash-flow sensitivity of rate-stressed households.", bold_prefix="GM implication")

    styled_heading(doc, "Consumer Confidence", level=2)
    body_para(doc, "ANZ-Roy Morgan Consumer Confidence Index: 67.2 as of 5 May 2026. This is near historic lows. Only 15% of Australians consider now a good time to buy major household goods. Pessimistic 5-year economic expectations rose 2 percentage points to 31%. Foot traffic and conversion assumptions built on pre-2025 benchmarks should be revised downward.")
    bullet_para(doc, "Source: Roy Morgan/ANZ, 5 May 2026", bold_prefix="Source")

    styled_heading(doc, "Household Spending Indicator", level=2)
    body_para(doc, "ABS Monthly Household Spending Indicator recorded Furnishings and Household Equipment +1.6% MoM in March 2026 (seasonally adjusted, current prices). Treat this figure with caution: it is nominal only and real volume growth is unconfirmed.")
    bullet_para(doc, "Nominal vs real volume unconfirmed. Do not use in GM presentations or board papers without ABS clarification.", bold_prefix="Review Flag", flag=True)
    bullet_para(doc, "Source: ABS MHSI March 2026", bold_prefix="Source")

    styled_heading(doc, "Housing Turnover", level=2)
    body_para(doc, "Approximately 559,457 residential sales recorded YTD 2026, down 1.9% YoY but 5.6% above the 5-year average. The market is two-speed: Perth and Queensland remain strong; Sydney and Melbourne are softening. Median dwelling value AUD $922,838 (+9.9% YoY). This bifurcation should inform franchise expansion prioritisation and stock positioning by corridor.")
    bullet_para(doc, "Source: Cotality April 2026", bold_prefix="Source")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 3: Sector Events
    # ------------------------------------------------------------------
    styled_heading(doc, "3. Sector Events", level=1)

    styled_heading(doc, "A.H. Beard — Voluntary Administration", level=2)
    body_para(doc, "A.H. Beard — Australia's largest domestic mattress manufacturer (126 years operating) — entered voluntary administration in late April 2026. Administrators P.A. Lucas & Co. confirm trading continues. Retail partners currently affected: Harvey Norman, Forty Winks, Domayne. No buyer has been named.", bold=False)
    body_para(doc, "Why this matters operationally:")
    bullet_para(doc, "Retail partners (especially Forty Winks) may face stock replenishment uncertainty, creating short-term range gaps Snooze can exploit in conversations with in-store staff and customers.")
    bullet_para(doc, "Competitor buyers (Harvey Norman, Nick Scali, Temple and Webster) may attempt to acquire A.H. Beard assets or contracts. Snooze/LKG should monitor for any exclusivity announcements.")
    bullet_para(doc, "If Snooze or G&G has any A.H. Beard supply relationship, confirm continuity with administrators within 48 hours.")
    action_box(doc, "Action Required", "Commercial team to confirm LKG supply chain exposure within 48 hours. Monitor for buyer announcement. Assess range gap opportunity in affected retailer catchments.")

    add_horizontal_rule(doc)

    styled_heading(doc, "ACCC — Emma Sleep $15M Penalty", level=2)
    body_para(doc, "The Federal Court ruled on 24 April 2026 that Emma Sleep's countdown timers and strikethrough pricing across 74 products constituted a 'deliberate marketing strategy' of misleading conduct. Over 4 million consumers received misleading emails; 500,000 received misleading SMS. ACCC penalty: AUD $15 million.")
    body_para(doc, "Why this matters operationally:")
    bullet_para(doc, "This is the enforcement action the ACCC has been signalling since its 2024 online pricing review. DTC mattress brands are now on notice.")
    bullet_para(doc, "Snooze operates countdown timers, flash sale mechanics, and strikethrough pricing across its e-commerce and email channels. These must be audited against the Federal Court's specific findings before the next promotional campaign.")
    bullet_para(doc, "Opportunity: position Snooze as the transparent, compliant alternative to DTC brands. 'Genuine advice, genuine prices' positioning is available and defensible.")
    action_box(doc, "Action Required", "Legal/Compliance to audit all active Snooze digital promotional mechanics against Emma Sleep ruling criteria. Complete before 14 May. Flag any non-compliant mechanics to GM.")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 4: Competitive Activity
    # ------------------------------------------------------------------
    styled_heading(doc, "4. Competitive Activity This Week", level=1)

    styled_heading(doc, "Active Promotions — Direct Competitors", level=2)

    comp_data = [
        ("Forty Winks", "Forty Frenzy", "Up to 50% off mattresses, 30% off bedroom furniture. Using bedMATCH diagnostic as in-store conversion tool.", "High", "Direct Snooze competitor. Simultaneous with Harvey Norman. Margin pressure is two-fronted."),
        ("Harvey Norman", "May 5-25 Catalogue", "Bedroom and mattress promotions live. Share buyback extended (~$900M) signals financial confidence.", "High", "Largest floor space competitor. Buyback signals HVN is not in retreat."),
        ("Koala", "EOFY Pre-Stage", "Up to 30% off mattresses + 50% off duvet bundle + $50K Volvo/cash giveaway.", "Medium-High", "High coupon-aggregator discoverability. Motivated online buyers being captured."),
        ("Sleeping Duck", "SLEEPSYSTEM2026", "$350 off bed+mattress bundle via promo code.", "Medium", "Code sourced via aggregator — verify on sleepingduck.com before using in competitive response briefings."),
        ("Ecosa", "Kids EOFY Pre-Stage", "30-35% off kids' mattresses and toppers.", "Medium", "Kids category signal — assess Future Sleep range alignment."),
        ("Fantastic Furniture", "Autumn Sale", "15-30% off bedroom packages.", "Medium", "Now Allegro-owned. New PE capital may intensify pricing. Watch for strategy shift post-acquisition."),
        ("IKEA AU", "PS 2026 Launch", "Launching 14 May. IKEA Family early access from ~9 May.", "High", "Mid-month launch will reset consumer attention. Snooze activation before 14 May advantageous."),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_header_row(table, ["Competitor", "Promotion", "Mechanics", "Confidence", "Snooze Implication"])

    for row_data in comp_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            para = cell.paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(9)
            if i == 3 and "Medium" in val and "High" not in val:
                run.font.color.rgb = AMBER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_paragraph()

    styled_heading(doc, "Strategic Competitor Moves", level=2)
    bullet_para(doc, "Nick Scali H1 FY26: ANZ revenue +13.1%, NPAT +29%, record share price. Planning 5 new stores in FY26; NZ potentially doubling to 13 stores. Nick Scali is the benchmark for margin-disciplined growth in this environment.", bold_prefix="Nick Scali (NCK)")
    bullet_para(doc, "Allegro Funds acquired Fantastic Furniture (86 stores, AUD $563.5M revenue) on 4 May 2026. PE-backed with 12-36 month investment horizon. Store overlap analysis with Snooze and G&G catchments is required.", bold_prefix="Fantastic Furniture acquisition")
    bullet_para(doc, "Combined Amart+Freedom entity (~126 stores, projected AUD $1B+ revenue — unverified) is eyeing a potential ASX IPO in 2026. Quadrant PE-backed. If IPO proceeds, this group will have access to public capital markets for further expansion.", bold_prefix="Amart + Freedom")
    bullet_para(doc, "IPO revenue figure of $1B+ is unverified. Do not use in board materials without independent confirmation.", bold_prefix="Review Flag", flag=True)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 5: Peer Performance Signals
    # ------------------------------------------------------------------
    styled_heading(doc, "5. Peer Performance Signals", level=1)

    body_para(doc, "The most important comparative signal this period is Nick Scali vs Adairs — two companies in the same consumer environment with opposite outcomes:")

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    add_table_header_row(table2, ["Company", "Revenue", "NPAT", "Gross Margin"])
    peers = [
        ("Nick Scali (NCK)", "+13.1% ANZ", "+29% to undisclosed", "Disciplined — no excessive promotional clearance"),
        ("Adairs (ADH)", "+5.9% to AUD $329M", "-34% to AUD $12.8M", "-170bps to 17.7% (Black Friday / Christmas clearance)"),
        ("Harvey Norman (HVN)", "+4.8% AU franchisee", "+15.2% to AUD $321.9M", "Stable — property income buffer"),
        ("Temple & Webster (TPW)", "+20% to AUD $375.9M", "Loss-making (growth phase)", "Online structural share gain: 2.9% market share (record)"),
    ]
    for row_data in peers:
        row = table2.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()
    body_para(doc, "The structural lesson: Adairs chose promotional matching. Nick Scali chose margin discipline. Same environment, same quarter, NPAT outcome difference of 63 percentage points. This is the live evidence base for LKG's EOFY promotional approach.")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 6: Recommended Actions
    # ------------------------------------------------------------------
    styled_heading(doc, "6. Recommended Actions", level=1)

    styled_heading(doc, "High Priority — Immediate (by 14 May 2026)", level=2)

    # IG-1
    styled_heading(doc, "IG-1: Recalculate and Hard-Floor Promotional Margins", level=3)
    body_para(doc, "Evidence: Foam costs up approximately 20% (industry estimate — confirm contracted price). Ocean freight surcharge $400-900 per 40ft container on China-Australia routes (Drewry WCI, 30 Apr 2026). Adairs lost 170bps gross margin from equivalent promotional posture in the same period.")
    action_box(doc, "Action Required", "Finance and Ops to pull current COGS for top 20 promotional mattress SKUs and recalculate at current input costs. Remove or reprice any SKU below minimum GM floor before 14 May. Do not launch EOFY promotions without this check.")
    bullet_para(doc, "Foam cost estimate is industry-level, not contracted LKG rate. Confirm with procurement before briefing stores.", bold_prefix="Review Flag", flag=True)
    bullet_para(doc, "Owner: Finance / Operations. Deadline: 14 May 2026.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    # IG-6
    styled_heading(doc, "IG-6: Lock Q3 2026 Freight Bookings Immediately", level=3)
    body_para(doc, "Evidence: Ocean freight surcharges $400-900 per 40ft container (EpicSourcing; Drewry WCI 30 Apr 2026). Australia Post fuel surcharge increased from 4.8% to 12% from 23 April 2026.")
    action_box(doc, "Action Required", "Supply chain team to review and lock Q3 2026 freight bookings within 7 days. Confirm e-commerce fulfilment cost models have been updated to reflect the Australia Post surcharge before EOFY campaign budgets are finalised.")
    bullet_para(doc, "Owner: Supply Chain / Logistics. Deadline: 15 May 2026.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    # IG-2
    styled_heading(doc, "IG-2: Post-Mother's Day Service-Led Conversion Activation (15-31 May)", level=3)
    body_para(doc, "Evidence: All major competitors wind down current promotions between 14-25 May. EOFY promotional cycle not expected until June. Approximately 10-15 day window of reduced competitive promotional noise.")
    action_box(doc, "Action Required", "Brief Snooze marketing team by 9 May. Design service-led activation (extended sleep consultation, mattress trial extension, local franchise events). Target existing CRM member database. No headline discounting required.")
    bullet_para(doc, "Owner: Marketing / Franchise Operations. Brief by: 9 May 2026.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    styled_heading(doc, "High Priority — Complete by End of May (Both: GM + Board visibility)", level=2)

    # IG-3
    styled_heading(doc, "IG-3: Formal Minimum Gross Margin Floor Policy — All Portfolio Companies", level=3)
    body_para(doc, "Evidence: Nick Scali NPAT +29% vs Adairs NPAT -34% in the same period. The only structural difference was margin discipline. Without a formalised floor policy, franchise and wholesale promotional decisions default to competitive matching — systematic margin destruction through EOFY.")
    action_box(doc, "Action Required", "GM and CFO to draft minimum GM floor policy by category (mattresses, bedroom furniture, accessories) for Snooze, Future Sleep, and G&G. Use NCK vs ADH comparison as evidentiary anchor. Table at June Board meeting.")
    bullet_para(doc, "Owner: GM / CFO. Board submission: June 2026.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    # IG-4
    styled_heading(doc, "IG-4: Commission Competitive Positioning Brief on Allegro / Fantastic Furniture Acquisition", level=3)
    body_para(doc, "Evidence: Allegro Funds acquired Fantastic Furniture (86 stores, AUD $563.5M revenue) on 4 May 2026. Amart+Freedom merger (Aug 2025) created approximately 120-store entity. Both PE-backed with 12-36 month investment horizons.")
    action_box(doc, "Action Required", "GM to commission store-location overlap analysis: Fantastic Furniture vs Snooze and G&G catchments. Include price-tier overlap and Allegro's known retail acquisition playbook. Table at June Board meeting.")
    bullet_para(doc, "Owner: GM / Commercial Strategy. Deadline: June Board meeting.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    # IG-7
    styled_heading(doc, "IG-7: G&G Furniture USD Contract Exposure and FX Hedging Review", level=3)
    body_para(doc, "Evidence: AUD/USD at 0.7242 (+13% YoY, four-year high) as of May 2026.")
    action_box(doc, "Action Required", "G&G CFO to provide written summary: (1) proportion of USD-denominated contracts, (2) current hedging position and maturity dates, (3) whether AUD tailwind is reflected in FY27 cost planning. Review within two weeks.")
    bullet_para(doc, "FX rate sourced from market aggregator — verify current rate with treasury before acting.", bold_prefix="Review Flag", flag=True)
    bullet_para(doc, "Owner: G&G CFO. Deadline: 22 May 2026.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    # IG-9
    styled_heading(doc, "IG-9: Housing Demand Readiness Plan for 2026-27 Wave", level=3)
    body_para(doc, "Evidence: Dwelling commencements +11.6% YoY in Q3 2025 (ABS). Typical lag to furniture purchase: 9-15 months, implying demand wave late 2026 to Q2 2027. March 2026 approvals fell 10.5% — wave may moderate post-2027.")
    action_box(doc, "Action Required", "GM to develop readiness brief covering: (1) franchise network capacity in growth corridors (Perth, QLD); (2) Future Sleep range alignment with new-mover segment; (3) inventory and lead-time planning for H2 2026. Table at June Board meeting.")
    bullet_para(doc, "Owner: GM / Franchise Operations. Deadline: June Board meeting.", bold_prefix="Owner and Deadline")

    doc.add_paragraph()

    styled_heading(doc, "Medium Priority — Complete Within Two Weeks", level=2)

    # IG-10
    styled_heading(doc, "IG-10: Benchmark Snooze Consumer Finance Offer; Consider BNPL Mechanic for EOFY", level=3)
    body_para(doc, "Evidence: Only 15% of Australians consider now a good time to buy major household items. Competitors (Forty Winks, Amart) are using sleep-trial mechanics and loyalty offers to reduce purchase friction rather than pure discounting. Motivated buyers who are cash-flow-hesitant — not price-objecting — are being lost.")
    action_box(doc, "Action Required", "Commercial team to benchmark Snooze vs Forty Winks, Bedshed, Harvey Norman, and Amart on: interest-free term length; BNPL integration prominence; in-store finance presentation at point of sale. Identify gaps and propose an EOFY mechanic within two weeks.")
    bullet_para(doc, "Owner: Commercial / Marketing. Deadline: 22 May 2026.", bold_prefix="Owner and Deadline")
    bullet_para(doc, "BNPL provider terms and current Snooze finance offer details require internal confirmation before external benchmarking.", bold_prefix="Review Flag", flag=True)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 7: Review Flags
    # ------------------------------------------------------------------
    styled_heading(doc, "7. Review Flags Requiring Resolution Before Acting", level=1)

    body_para(doc, "The following items carry uncertainty that must be resolved internally before the associated action or figure is used in presentations, board papers, or external communications.", colour=AMBER)

    flags = [
        ("IG-1 / ABS MHSI", "ABS March 2026 +1.6% MoM figure is nominal only. Real volume unconfirmed. Do not use as evidence of category growth without ABS clarification.", "Finance / Research"),
        ("IG-1 / Foam costs", "~20% cost increase is an industry-level estimate. Confirm contracted LKG rate with Procurement before briefing stores or board.", "Procurement"),
        ("CA Amart+Freedom revenue", "AUD $1B+ revenue figure is unverified (Inside Retail / ChannelNews). Do not include in board papers or investor materials without independent confirmation.", "Commercial / Finance"),
        ("Sleeping Duck promo code", "SLEEPSYSTEM2026 code sourced via aggregator (BHG). Verify at sleepingduck.com before using in competitive response briefings.", "Marketing"),
        ("Tempur-AI partnership AU applicability", "US announcement only. AU activation timeline unconfirmed. Do not include in AU competitive threat analysis without confirmation.", "Research"),
        ("Comps valuation multiples", "EV/EBITDA ranges for NCK, HVN, TPW sourced from aggregators, not broker research. Do not use in M&A or capital allocation discussions without broker-verified data.", "Finance / CFO"),
        ("IG-7 / AUD/USD rate", "0.7242 rate sourced from market aggregator. Verify current rate and hedging position with G&G treasury before acting.", "G&G CFO / Treasury"),
    ]

    table3 = doc.add_table(rows=1, cols=3)
    table3.style = "Table Grid"
    add_table_header_row(table3, ["Item / Source", "Issue", "Owner"])
    for row_data in flags:
        row = table3.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = AMBER

    doc.add_paragraph()

    # Disclaimer
    disclaimer_footer(doc)

    out_path = os.path.join(OUTPUT_DIR, "lkg-furniture-gm-weekly-digest.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}  ({os.path.getsize(out_path):,} bytes)")
    return out_path


# ===========================================================================
# DOCUMENT 2 — BOARD WEEKLY DIGEST
# ===========================================================================

def build_board_digest():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc_cover(
        doc,
        "LKG Furniture — Board Weekly Digest",
        "Hypnos Group (Snooze | Future Sleep | G&G Furniture)",
        "LKG Board of Directors"
    )

    # ------------------------------------------------------------------
    # SECTION 1: Board Summary
    # ------------------------------------------------------------------
    styled_heading(doc, "1. Board Summary", level=1)
    body_para(doc, "Five strategic signals require Board-level awareness this week. Three require a formal response at the June Board meeting.", bold=True)

    board_bullets = [
        ("Rate cycle risk is not priced in", "RBA's third consecutive hike to 4.35% and consumer confidence at 67.2 (near 50-year lows) create material downside risk to FY27 revenue assumptions. The Board requires stress-tested revenue scenarios before capital allocation decisions for FY27 can be made."),
        ("A.H. Beard administration is a rare M&A and positioning event", "The voluntary administration of Australia's largest domestic mattress manufacturer creates a time-bounded window for asset acquisition, brand positioning, or supply chain advantage. Board should be briefed on LKG's strategic options within 30 days."),
        ("PE capital is restructuring the competitive landscape", "Allegro Funds acquired Fantastic Furniture (4 May); Quadrant PE's Amart+Freedom entity (~126 stores) is targeting an ASX IPO. The Board needs a competitive response framework before PE investment activity shows up in LKG's like-for-like data."),
        ("Margin discipline is the single most important operational lever", "Nick Scali NPAT +29%; Adairs NPAT -34% in the same quarter. The structural difference was margin discipline. A formal minimum GM floor policy across the LKG portfolio should be tabled at the June Board meeting."),
        ("Online structural shift is accelerating", "Temple and Webster achieved record 2.9% market share (+20% revenue) in the same environment where physical retailers declined. The Board needs to set the capital allocation threshold for LKG's online investment or the cost of deferral compounds each cycle."),
    ]
    for bp, bt in board_bullets:
        bullet_para(doc, bt, bold_prefix=bp)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 2: Macro Environment
    # ------------------------------------------------------------------
    styled_heading(doc, "2. Macro Environment and Consumer Outlook", level=1)

    styled_heading(doc, "Interest Rate Cycle", level=2)
    body_para(doc, "The RBA raised the cash rate to 4.35% on 5 May 2026, the third consecutive increase. CBA economists expect a hold for the remainder of 2026, but further increases remain possible. This is a direct and sustained headwind to discretionary household goods spending.")
    bullet_para(doc, "Source: RBA Media Release mr-26-12, 5 May 2026.", bold_prefix="Source")
    bullet_para(doc, "Three consecutive hikes with no near-term relief signal materially increases the probability of a revenue downside scenario through H1 FY27. Board should request stress-tested revenue scenarios.", bold_prefix="Strategic implication")

    styled_heading(doc, "Consumer Confidence — Near Historic Lows", level=2)
    body_para(doc, "ANZ-Roy Morgan Consumer Confidence: 67.2 (5 May 2026). Only 15% of Australians consider now a good time to buy major household goods. Pessimistic 5-year economic expectations rose 2 percentage points to 31%. This is the seventh lowest reading in 50 years of the index.")
    bullet_para(doc, "Source: Roy Morgan/ANZ, 5 May 2026.", bold_prefix="Source")
    bullet_para(doc, "Any LKG FY27 revenue plan built on a consumer recovery assumption requires explicit Board-level acknowledgement of this risk. Confidence at this level has historically preceded a 12-18 month period of suppressed discretionary spending.", bold_prefix="Strategic implication")

    styled_heading(doc, "Housing Market — Two-Speed Dynamics", level=2)
    body_para(doc, "Approximately 559,457 residential sales YTD 2026, down 1.9% YoY but 5.6% above the 5-year average. Median dwelling value AUD $922,838 (+9.9% YoY). Market is bifurcated: Perth and Queensland strong; Sydney and Melbourne softening. Dwelling commencements +11.6% YoY in Q3 2025 imply a furniture demand wave in late 2026 to Q2 2027.")
    bullet_para(doc, "Source: Cotality April 2026; ABS.", bold_prefix="Source")
    bullet_para(doc, "PE-backed competitors are likely preparing for the same demand wave. LKG's franchise network and inventory positioning in growth corridors (Perth, QLD) needs to be validated against this timeline.", bold_prefix="Strategic implication")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 3: Sector Events with Strategic Implications
    # ------------------------------------------------------------------
    styled_heading(doc, "3. Sector Events with Strategic Implications", level=1)

    styled_heading(doc, "A.H. Beard — Voluntary Administration", level=2)
    body_para(doc, "A.H. Beard — Australia's largest domestic mattress manufacturer, 126 years operating — entered voluntary administration in late April 2026. Administrators P.A. Lucas & Co. Retail distribution partners affected: Harvey Norman, Forty Winks, Domayne. No buyer named.")
    body_para(doc, "Strategic implications for LKG:")
    bullet_para(doc, "This is the most significant supply-side disruption in the Australian sleep sector in decades. The window before a buyer is named is the time for LKG to assess whether any A.H. Beard assets — brand, manufacturing relationships, or retailer contracts — are strategically relevant.")
    bullet_para(doc, "Competitor retailers (Harvey Norman, Nick Scali) may attempt to acquire assets to secure exclusive supply. Board should determine whether LKG wishes to be a bidder or a monitor.")
    bullet_para(doc, "Source: BedTimes Magazine, April 2026.", bold_prefix="Source")

    styled_heading(doc, "ACCC — Emma Sleep $15M Penalty", level=2)
    body_para(doc, "Federal Court ruling (24 April 2026): Emma Sleep's countdown timers and strikethrough pricing were a 'deliberate marketing strategy' of misleading conduct. ACCC penalty: AUD $15 million. 4 million+ consumers received misleading emails; 500,000 SMS messages.")
    body_para(doc, "Strategic implications for LKG:")
    bullet_para(doc, "This ruling elevates compliance risk for all Australian furniture and bedding retailers using dynamic promotional mechanics. LKG's compliance exposure should be reviewed by legal counsel.")
    bullet_para(doc, "Reputational opportunity: LKG can differentiate its brands on transparent, compliant promotional practices at a time when a major DTC competitor has been publicly penalised. This is a brand positioning opportunity, not merely a compliance risk.")
    bullet_para(doc, "Source: ACCC Media Release, April 2026.", bold_prefix="Source")

    styled_heading(doc, "PE Capital Entering the Sector", level=2)
    body_para(doc, "Two concurrent PE moves are reshaping the competitive landscape:")
    bullet_para(doc, "Allegro Funds acquired Fantastic Furniture (86 stores, AUD $563.5M revenue) on 4 May 2026.", bold_prefix="Allegro / Fantastic Furniture")
    bullet_para(doc, "Quadrant PE-backed combined Amart+Freedom entity (~126 stores, projected AUD $1B+ revenue — unverified) is targeting an ASX IPO in 2026.", bold_prefix="Amart + Freedom IPO")
    body_para(doc, "Both entities have PE capital backing with 12-36 month investment horizons and explicit expansion targets. PE-backed retailers typically prioritise store rollout, system investment, and promotional aggression in the growth phase.")
    bullet_para(doc, "Board should assess whether LKG's current competitive positioning and capital allocation is adequate to defend against two simultaneously-capitalised competitors. A formal competitive response plan should be tabled at the June Board meeting.", bold_prefix="Strategic implication")
    bullet_para(doc, "Amart+Freedom $1B+ revenue figure is unverified. Do not use in public-facing materials.", bold_prefix="Review Flag", flag=True)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 4: Peer Performance Benchmarks
    # ------------------------------------------------------------------
    styled_heading(doc, "4. Peer Performance Benchmarks", level=1)

    body_para(doc, "All figures are H1 FY26 unless otherwise stated. Valuation multiples are directional only — sourced from aggregators, not broker research.")
    bullet_para(doc, "Valuation multiples below are directional only. Aggregator-sourced, not broker-verified. Do not use in M&A or capital allocation analysis without independent verification.", bold_prefix="Review Flag", flag=True)

    doc.add_paragraph()

    table4 = doc.add_table(rows=1, cols=6)
    table4.style = "Table Grid"
    add_table_header_row(table4, ["Company", "Revenue", "NPAT", "Gross Margin", "EV/EBITDA (indicative)", "Board Relevance"])
    peers_b = [
        ("Nick Scali (NCK)", "+13.1% ANZ", "+29%", "Disciplined — no clearance", "~8.7-9.2x", "Closest pure-play comp. Margin discipline is the differentiator."),
        ("Adairs (ADH)", "+5.9% to $329M", "-34% to $12.8M", "-170bps to 17.7%", "Not directly applicable", "Risk case: promotional matching destroys NPAT in a single half."),
        ("Harvey Norman (HVN)", "+4.8% AU", "+15.2% to $321.9M", "Stable", "~8-10x (distorted by property)", "Competitor. Property buffer insulates margin — LKG does not have this."),
        ("Temple & Webster (TPW)", "+20% to $375.9M", "Loss-making (growth)", "Online premium", "~27.2x (growth premium)", "Online structural shift signal. 2.9% market share — record."),
        ("LKG implied private range", "—", "—", "—", "~6.0-7.5x (20-30% illiquidity discount vs NCK)", "Reference range only. Board should confirm with advisers before any transaction."),
    ]
    for row_data in peers_b:
        row = table4.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()
    body_para(doc, "The Nick Scali vs Adairs comparison is the live same-period benchmark. NCK's NPAT +29% vs ADH's -34% in the same consumer environment. The structural differentiator was margin discipline. This is the evidentiary anchor for the proposed minimum GM floor policy.")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 5: Competitive Landscape Shifts
    # ------------------------------------------------------------------
    styled_heading(doc, "5. Competitive Landscape Shifts", level=1)

    body_para(doc, "Three structural shifts are occurring simultaneously:")

    bullet_para(doc, "PE-backed competitors (Allegro/Fantastic Furniture; Quadrant/Amart+Freedom) are being recapitalised and expanding. Both have stated store growth targets and investment horizons. This is not a cyclical competitive shift — it is structural.", bold_prefix="PE capital consolidation")
    bullet_para(doc, "Temple and Webster achieved record 2.9% online market share in H1 FY26, growing at +20% in a flat-to-declining physical retail environment. The online shift is accelerating, not plateauing.", bold_prefix="Online structural shift")
    bullet_para(doc, "Nick Scali is planning 5 new stores in FY26 and potentially doubling its NZ presence to 13 stores. A well-capitalised, margin-disciplined competitor is actively expanding into growth corridors.", bold_prefix="Nick Scali expansion")

    body_para(doc, "Each of these three shifts individually would warrant a Board response. Occurring simultaneously, they represent a step-change in the competitive intensity facing LKG's portfolio companies.")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 6: Strategic Recommendations for Board
    # ------------------------------------------------------------------
    styled_heading(doc, "6. Strategic Recommendations for Board", level=1)

    styled_heading(doc, "IG-5: Revenue Scenario Framework for Rate-Cycle Downside", level=2)
    body_para(doc, "Evidence: Three consecutive RBA hikes to 4.35% (5 May 2026). Consumer confidence 67.2 — near historic lows. Harvey Norman Australian LFL was -1.2% at peak Christmas 2025.")
    body_para(doc, "Recommended action: Portfolio CFOs and GM to prepare three scenarios — (1) Base: rate pause Q3, confidence recovery H2; (2) Upside: rate cut, housing demand wave; (3) Downside: two further hikes, confidence falls below 65. Each scenario to include revenue and GM implications per business unit.")
    bullet_para(doc, "Table at June Board meeting.", bold_prefix="Timeline")
    bullet_para(doc, "Board cannot make FY27 capital allocation decisions without stress-tested assumptions. This is a prerequisite for all other strategic planning.", bold_prefix="Why it matters")

    doc.add_paragraph()

    styled_heading(doc, "IG-3: Minimum Gross Margin Floor Policy — All Portfolio Companies", level=2)
    body_para(doc, "Evidence: Nick Scali NPAT +29% vs Adairs NPAT -34% in H1 FY26. The only structural difference was margin discipline. Without a formalised floor, franchise and wholesale promotional decisions default to competitive matching.")
    body_para(doc, "Recommended action: GM and CFO to draft minimum GM floor by category (mattresses, bedroom furniture, accessories) for Snooze, Future Sleep, and G&G. Use NCK vs ADH as evidentiary anchor.")
    bullet_para(doc, "Table at June Board meeting.", bold_prefix="Timeline")
    bullet_para(doc, "Without a formalised floor, EOFY will systematically erode margin across the portfolio. One poor half can eliminate a year of revenue growth.", bold_prefix="Why it matters")

    doc.add_paragraph()

    styled_heading(doc, "IG-4: Competitive Response to Allegro / Fantastic Furniture Acquisition", level=2)
    body_para(doc, "Evidence: Allegro acquired Fantastic Furniture (4 May 2026). Amart+Freedom merger Aug 2025. Both PE-backed, both expanding. Combined they represent a new tier of PE-capitalised competition below the Harvey Norman level but above the independent operator level — directly competing with Snooze and G&G.")
    body_para(doc, "Recommended action: GM to commission store-location overlap analysis and Allegro acquisition playbook brief. Board to consider whether a formal competitive response plan is required.")
    bullet_para(doc, "Table at June Board meeting.", bold_prefix="Timeline")
    bullet_para(doc, "PE capital entering the sector simultaneously changes the rules of engagement. Without a pre-planned response, LKG will react to competitor moves rather than anticipating them.", bold_prefix="Why it matters")

    doc.add_paragraph()

    styled_heading(doc, "IG-8: Online Investment Threshold — Benchmark vs Temple and Webster Trajectory", level=2)
    body_para(doc, "Evidence: Temple and Webster H1 FY26 revenue +20% to $376M; 2.9% market share (record). This was achieved in the same weak consumer environment where physical retailers declined.")
    body_para(doc, "Recommended action: Analyst to benchmark LKG portfolio online revenue share vs Temple and Webster trajectory. Present capital allocation framing for what investment level would arrest (not reverse) structural share shift.")
    bullet_para(doc, "Table at Q3 FY26 Board meeting.", bold_prefix="Timeline")
    bullet_para(doc, "Online share shift is structural, not cyclical. Every cycle without a clear online investment threshold allows the gap to widen. Deferral makes the eventual investment larger and the catch-up less certain.", bold_prefix="Why it matters")

    doc.add_paragraph()

    styled_heading(doc, "IG-7: G&G Furniture USD Contract Exposure and FX Hedging Review", level=2)
    body_para(doc, "Evidence: AUD/USD at 0.7242 (+13% YoY, four-year high) as of May 2026.")
    body_para(doc, "Recommended action: G&G CFO to provide Board with written summary of USD-denominated contract proportion, current hedging position and maturity dates, and whether the AUD tailwind is reflected in FY27 cost planning.")
    bullet_para(doc, "Within two weeks. Written summary to Board.", bold_prefix="Timeline")
    bullet_para(doc, "13% AUD appreciation on USD contracts is a significant cost tailwind — or an unrecognised hedging risk if contracts were locked at worse rates. This must be confirmed before FY27 budgets are finalised.", bold_prefix="Why it matters")
    bullet_para(doc, "AUD/USD rate sourced from market aggregator. Verify with G&G treasury.", bold_prefix="Review Flag", flag=True)

    doc.add_paragraph()

    styled_heading(doc, "IG-9: Housing Demand Readiness Plan for 2026-27 Wave", level=2)
    body_para(doc, "Evidence: Dwelling commencements +11.6% YoY in Q3 2025 implies a furniture demand wave in late 2026 to Q2 2027. Perth and QLD are leading growth corridors. PE-backed competitors are likely preparing for the same wave.")
    body_para(doc, "Recommended action: GM to develop readiness brief covering franchise network capacity in growth corridors, Future Sleep range alignment with new-mover segment, and inventory/lead-time planning for H2 2026.")
    bullet_para(doc, "Table at June Board meeting.", bold_prefix="Timeline")
    bullet_para(doc, "This is a time-bounded structural demand opportunity. The wave arrives regardless of consumer confidence. LKG needs to be positioned to capture it, not react to it after competitors have claimed the corridor.", bold_prefix="Why it matters")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 7: Review Flags
    # ------------------------------------------------------------------
    styled_heading(doc, "7. Review Flags Requiring Board Attention", level=1)

    body_para(doc, "The following data points underpin Board-level recommendations but carry uncertainty that must be resolved before being cited in formal Board papers or external communications.", colour=AMBER)

    board_flags = [
        ("Comps valuation multiples", "EV/EBITDA ranges are aggregator-sourced, not broker-verified. LKG implied range of 6.0-7.5x is directional only. Do not use in M&A or capital allocation discussions without independent verification.", "CFO / Corporate Adviser"),
        ("ABS spending figure", "+1.6% MoM nominal only — real volume unconfirmed. Do not use as evidence of category growth.", "Finance / Research"),
        ("Amart+Freedom revenue", "AUD $1B+ unverified. Do not include in Board papers without independent confirmation.", "Commercial / Finance"),
        ("AUD/USD rate", "0.7242 from market aggregator. Verify with G&G treasury before acting.", "G&G CFO / Treasury"),
        ("A.H. Beard M&A options", "No buyer named. Administrators are P.A. Lucas & Co. Board should determine whether to formally monitor or engage within 30 days.", "Board / M&A Adviser"),
    ]

    table5 = doc.add_table(rows=1, cols=3)
    table5.style = "Table Grid"
    add_table_header_row(table5, ["Item", "Issue", "Owner"])
    for row_data in board_flags:
        row = table5.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = AMBER

    doc.add_paragraph()

    disclaimer_footer(doc)

    out_path = os.path.join(OUTPUT_DIR, "lkg-furniture-board-weekly-digest.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}  ({os.path.getsize(out_path):,} bytes)")
    return out_path


# ===========================================================================
# DOCUMENT 3 — INTERNAL SOURCE LOG
# ===========================================================================

def build_source_log():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    doc_cover(
        doc,
        "LKG Furniture — Internal Source Log",
        "Full Item Registry | All Classifications | All Sources",
        "Internal Research Team — Not for Distribution"
    )

    body_para(doc, "This document is the complete research log for the week of May 1-8, 2026. It includes all items reviewed, their classification, confidence rating, source, review flags, and the human approval decision. It is the audit record for the GM and Board digests produced from this week's research cycle.", colour=DARK_GREY)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 1: Full Item Log
    # ------------------------------------------------------------------
    styled_heading(doc, "1. Full Item Log", level=1)

    all_items = [
        # ID | Summary | Classification | Confidence | Review Flag | Source | Date
        ("SO-1", "RBA raises cash rate to 4.35% — third consecutive hike", "Board / Both", "Very High", "", "RBA Media Release mr-26-12", "5 May 2026"),
        ("SO-2", "ANZ-Roy Morgan Consumer Confidence 67.2 — near historic lows; 15% consider now good time to buy household goods", "Both", "Very High", "", "Roy Morgan/ANZ", "5 May 2026"),
        ("SO-3", "ABS MHSI: Furnishings +1.6% MoM March 2026 — nominal only", "Both", "Medium", "Nominal vs real unconfirmed", "ABS Monthly Household Spending Indicator, March 2026", "April 2026"),
        ("SO-4", "Housing turnover 559,457 YTD -1.9% YoY but +5.6% vs 5yr avg; median $922,838 +9.9% YoY; Perth/QLD strong, Syd/Melb soft", "Both / Board", "High", "", "Cotality April 2026", "April 2026"),
        ("SO-5", "A.H. Beard voluntary administration — P.A. Lucas & Co; continues trading; Harvey Norman, Forty Winks, Domayne affected", "Both / Board", "Very High", "No buyer named — monitor daily", "BedTimes Magazine", "April 2026"),
        ("SO-6", "ACCC $15M penalty vs Emma Sleep — countdown timers, strikethrough pricing, 74 products, 4M+ misleading emails", "Both / Board", "Very High", "", "ACCC Media Release", "24 April 2026"),
        ("SO-7", "Allegro Funds acquires Fantastic Furniture (86 stores, AUD $563.5M revenue) 4 May 2026", "Board / Both", "Very High", "", "Inside Retail; Allegro Funds announcement", "4 May 2026"),
        ("SO-8", "Nick Scali H1 FY26: ANZ revenue +13.1%, NPAT +29%, record share price; 5 new stores FY26 planned", "Board / Both", "Very High", "", "Investing.com; ASX filings (NCK)", "May 2026"),
        ("SO-9", "Adairs H1 FY26: Revenue +5.9% to $329M, NPAT -34% to $12.8M, GM -170bps to 17.7%", "Board / Both", "Very High", "", "Kalkine; Yahoo Finance; ASX filings (ADH)", "May 2026"),
        ("CA-1", "Forty Winks 'Forty Frenzy' sale: up to 50% off mattresses, 30% off bedroom; bedMATCH diagnostic in-store", "GM", "High", "", "fortywinks.com.au", "Week of 5 May 2026"),
        ("CA-2", "Harvey Norman May 5-25 catalogue: bedroom and mattress promos; share buyback extended to Nov 2026 (~$900M)", "GM / Both", "High", "", "ASX trading updates; Kimbino", "Week of 5 May 2026"),
        ("CA-3", "Harvey Norman H1 FY26: after-tax profit $321.9M +15.2%, AU franchisee sales $3.5B +4.8%", "Board", "Very High", "", "Inside Retail; Motley Fool", "May 2026"),
        ("CA-4", "IKEA PS 2026 collection launching 14 May AU; IKEA Family early access", "GM", "High", "", "IKEA AU website", "Week of 5 May 2026"),
        ("CA-5", "Koala: up to 30% off mattresses + 50% off duvet + $50K Volvo/cash giveaway", "GM", "Medium-High", "", "Koala AU; Better Homes and Gardens aggregator", "Week of 5 May 2026"),
        ("CA-6", "Duplicate entry — removed", "Ignored", "N/A", "Duplicate — removed from research", "N/A", "N/A"),
        ("CA-7", "Sleeping Duck: SLEEPSYSTEM2026 code, $350 off bed+mattress bundle", "GM", "Medium", "Verify at sleepingduck.com — aggregator sourced", "BHG aggregator", "Week of 5 May 2026"),
        ("CA-8", "Ecosa: 30-35% off kids' mattresses and toppers; EOFY pre-marketing", "GM", "Medium", "", "Don't Pay Full aggregator", "Week of 5 May 2026"),
        ("CA-9", "Fantastic Furniture Autumn Sale: 15-30% off bedroom packages (now Allegro-owned)", "GM", "Medium", "", "BHG aggregator", "Week of 5 May 2026"),
        ("CA-10", "Super Retail Group LFL +0.4% — no furniture exposure", "Ignored", "High", "No LKG overlap", "ASX trading update (SUL)", "6 May 2026"),
        ("CA-11", "Amart+Freedom combined group (~126 stores, projected $1B+ rev — unverified) eyes ASX IPO 2026; Quadrant PE-backed", "Board", "Medium", "Revenue figure unverified", "Inside Retail; ChannelNews", "May 2026"),
        ("CA-12", "Nick Scali NZ expansion: potentially doubling to 13 stores FY26", "GM", "Very High", "", "ASX filings (NCK)", "May 2026"),
        ("CMP-1", "Super Retail Group trading update — no furniture exposure", "Ignored", "High", "No LKG category overlap", "ASX (SUL)", "6 May 2026"),
        ("CMP-2", "Nick Scali EV/EBITDA ~8.7-9.2x (directional, aggregator)", "Board", "Medium", "Aggregator-sourced — not broker-verified", "Market aggregators", "May 2026"),
        ("CMP-3", "Harvey Norman EV/EBITDA ~8-10x (distorted by property assets)", "Board", "Medium", "Property distortion; aggregator-sourced", "Market aggregators", "May 2026"),
        ("CMP-4", "Temple & Webster EV/EBITDA ~27.2x (online growth premium)", "Board", "Medium", "Aggregator-sourced — not broker-verified", "Market aggregators", "May 2026"),
        ("CMP-5", "LKG implied private reference range: ~6.0-7.5x EV/EBITDA (20-30% illiquidity discount vs NCK)", "Board", "Low-Medium", "Directional only. Requires independent verification before any transaction.", "Internal derivation from aggregator comps", "May 2026"),
        ("CMP-6", "Temple & Webster analyst fair value cut to $7.30 (4 May); +5.6% move 7 May (catalyst unconfirmed)", "Board", "Medium", "7 May catalyst unconfirmed", "Multiple analyst sources; ASX", "4-7 May 2026"),
        ("IG-1", "Recalculate and hard-floor promotional margins before 14 May — foam up ~20%, freight +$400-900/40ft", "GM", "High", "Foam cost is industry estimate; confirm contracted LKG rate", "Drewry WCI; EpicSourcing; Adairs H1 FY26", "May 2026"),
        ("IG-2", "Post-Mother's Day service-led conversion activation 15-31 May — brief marketing by 9 May", "GM", "High", "", "Competitive calendar analysis", "May 2026"),
        ("IG-3", "Formal minimum GM floor policy for all portfolio companies before EOFY — table June Board", "Both", "Very High", "", "NCK vs ADH H1 FY26 comparison", "May 2026"),
        ("IG-4", "Commission competitive positioning brief on Allegro/Fantastic Furniture acquisition — table June Board", "Both", "Very High", "", "Inside Retail; Allegro Funds announcement", "May 2026"),
        ("IG-5", "Revenue scenario framework for rate-cycle downside — three scenarios for June Board", "Board", "Very High", "", "RBA; ANZ-Roy Morgan; HVN LFL data", "May 2026"),
        ("IG-6", "Lock Q3 2026 freight bookings immediately — ocean freight and AusPost surcharges live", "GM", "High", "", "Drewry WCI; EpicSourcing; Australia Post", "30 Apr - 5 May 2026"),
        ("IG-7", "G&G USD contract exposure and FX hedging review — AUD/USD +13% YoY at 0.7242", "Both", "Medium", "AUD/USD rate from aggregator — verify with treasury", "Market aggregators", "May 2026"),
        ("IG-8", "Online investment threshold — benchmark vs Temple & Webster trajectory; capital allocation framing for Board", "Board", "High", "", "TPW H1 FY26 filings; market share data", "May 2026"),
        ("IG-9", "Housing demand readiness plan 2026-27 wave — franchise capacity, Future Sleep alignment, inventory planning", "Both", "High", "", "ABS dwelling commencements; Cotality", "May 2026"),
        ("IG-10", "Benchmark Snooze consumer finance offer; consider BNPL mechanic for EOFY — within two weeks", "GM", "Medium", "Current Snooze finance terms require internal confirmation", "Competitor website review; ANZ-Roy Morgan", "May 2026"),
    ]

    table_log = doc.add_table(rows=1, cols=7)
    table_log.style = "Table Grid"
    add_table_header_row(table_log, ["ID", "Summary", "Classification", "Confidence", "Review Flag", "Source", "Date"])

    col_widths = [Cm(1.2), Cm(6.5), Cm(2.0), Cm(1.8), Cm(3.5), Cm(4.0), Cm(2.0)]
    for i, width in enumerate(col_widths):
        for cell in table_log.columns[i].cells:
            cell.width = width

    for item in all_items:
        row = table_log.add_row()
        for i, val in enumerate(item):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i == 0:
                run.bold = True
                if item[2] == "Ignored":
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                elif "Board" in item[2]:
                    run.font.color.rgb = DARK_NAVY
                elif "GM" in item[2]:
                    run.font.color.rgb = MID_BLUE
                elif "Both" in item[2]:
                    run.font.color.rgb = RGBColor(0x37, 0x86, 0x60)
            if i == 4 and val:
                run.font.color.rgb = AMBER
            if item[2] == "Ignored":
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 2: Ignored Items
    # ------------------------------------------------------------------
    styled_heading(doc, "2. Ignored Items — Reasons", level=1)

    ignored = [
        ("CA-6", "Duplicate entry", "Identified as duplicate during research compilation. No unique content. Removed from all downstream documents."),
        ("CA-10", "Super Retail Group LFL +0.4%", "No furniture or bedding category exposure. SUL operates sporting goods, auto, and camping. No LKG overlap. Filed but not actioned."),
        ("CMP-1", "Super Retail Group trading update", "Same as CA-10. No category overlap. Filed for completeness."),
    ]

    table_ign = doc.add_table(rows=1, cols=3)
    table_ign.style = "Table Grid"
    add_table_header_row(table_ign, ["ID", "Item", "Reason for Ignoring"])
    for row_data in ignored:
        row = table_ign.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 3: Review Flags Detail
    # ------------------------------------------------------------------
    styled_heading(doc, "3. Review Flags — Detail and Resolution Required", level=1)

    body_para(doc, "The following flags were raised during research. Each must be resolved by the named owner before the associated item is used in formal Board papers, investor materials, or external communications.", colour=AMBER)

    flags_detail = [
        ("SO-3 / ABS MHSI", "Finance / Research", "+1.6% MoM is nominal only. Real volume growth unconfirmed. Do not cite as evidence of real category growth.", "Open", "Obtain ABS footnote or clarification note before next Board cycle."),
        ("IG-1 / Foam cost", "Procurement", "~20% cost increase is industry-level estimate from trade press. LKG contracted rate may differ significantly.", "Open", "Procurement to confirm contracted foam cost within 5 business days."),
        ("CA-11 / Amart+Freedom revenue", "Commercial / Finance", "$1B+ projected revenue is unverified (trade press). No ASX filing or audited source.", "Open", "Do not use in Board papers. Seek independent confirmation from ASIC/APRA registrations."),
        ("CA-7 / Sleeping Duck promo", "Marketing", "SLEEPSYSTEM2026 code sourced via BHG aggregator. May be expired or incorrect.", "Open", "Marketing to verify at sleepingduck.com before citing in competitive briefs."),
        ("Tempur-AI / AU applicability", "Research", "US announcement only. AU activation timeline unconfirmed.", "Open", "Do not include in AU competitive threat analysis. Monitor Tempur AU press releases."),
        ("CMP-2 through CMP-5 / Valuation multiples", "CFO / Corporate Adviser", "All EV/EBITDA figures are aggregator-sourced. LKG implied range (6.0-7.5x) is an internal derivation.", "Open", "Obtain broker research (at least one verified source) before any transaction or investor conversation."),
        ("IG-7 / AUD/USD rate", "G&G CFO / Treasury", "0.7242 sourced from market aggregator on approximate date. Hedging position unknown.", "Open", "G&G CFO to confirm current rate, hedging book, and FY27 cost plan within 2 weeks."),
        ("CMP-6 / TPW 7 May catalyst", "Research", "+5.6% single-day move on 7 May — catalyst unconfirmed.", "Open", "Research to identify catalyst from ASX announcements or analyst notes."),
    ]

    table_flags = doc.add_table(rows=1, cols=5)
    table_flags.style = "Table Grid"
    add_table_header_row(table_flags, ["Item / Source", "Owner", "Issue", "Status", "Resolution Path"])
    for row_data in flags_detail:
        row = table_flags.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i == 0:
                run.bold = True
                run.font.color.rgb = AMBER
            if i == 3:
                run.font.color.rgb = ACTION_RED

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # SECTION 4: Source List
    # ------------------------------------------------------------------
    styled_heading(doc, "4. Source List — All Sources Cited This Week", level=1)

    sources = [
        ("RBA Media Release mr-26-12", "5 May 2026", "rba.gov.au/media-releases/2026/mr-26-12.html", "Very High", "SO-1, IG-5"),
        ("ANZ-Roy Morgan Consumer Confidence", "5 May 2026", "roymorgan.com (weekly release)", "Very High", "SO-2, IG-10"),
        ("ABS Monthly Household Spending Indicator — March 2026", "April 2026", "abs.gov.au/statistics/economy/finance/monthly-household-spending-indicator", "Medium (nominal only)", "SO-3"),
        ("Cotality Residential Property Report — April 2026", "April 2026", "cotality.com.au", "High", "SO-4, IG-9"),
        ("BedTimes Magazine — A.H. Beard Administration", "April 2026", "bedtimesmagazine.com", "Very High", "SO-5, IG-4"),
        ("ACCC Media Release — Emma Sleep Penalty", "24 April 2026", "accc.gov.au/media-release/emma-sleep (indicative URL)", "Very High", "SO-6"),
        ("Inside Retail — Allegro Funds / Fantastic Furniture", "4-5 May 2026", "insideretail.com.au", "Very High", "SO-7, IG-4"),
        ("Inside Retail / ChannelNews — Amart+Freedom IPO", "May 2026", "insideretail.com.au; channelnews.com.au", "Medium (unverified)", "CA-11"),
        ("ASX Filings — Nick Scali (NCK)", "May 2026", "asx.com.au/asx/share-price-research/company/NCK", "Very High", "SO-8, CA-12, CMP-2"),
        ("Kalkine / Yahoo Finance — Adairs (ADH)", "May 2026", "kalkinemedia.com; finance.yahoo.com", "Very High", "SO-9, CMP-4"),
        ("ASX Filings / Kimbino — Harvey Norman (HVN)", "May 2026", "asx.com.au/asx/share-price-research/company/HVN; kimbino.com.au", "Very High", "CA-2, CA-3, CMP-3"),
        ("Forty Winks — fortywinks.com.au", "Week of 5 May 2026", "fortywinks.com.au", "High", "CA-1"),
        ("IKEA AU — IKEA.com/au", "Week of 5 May 2026", "ikea.com/au", "High", "CA-4"),
        ("Koala AU; Better Homes and Gardens aggregator", "Week of 5 May 2026", "au.koala.com; bhg.com.au", "Medium-High", "CA-5"),
        ("BHG aggregator — Sleeping Duck", "Week of 5 May 2026", "bhg.com.au (aggregator — unverified)", "Medium", "CA-7"),
        ("Don't Pay Full aggregator — Ecosa", "Week of 5 May 2026", "dontpayfull.com (aggregator)", "Medium", "CA-8"),
        ("BHG aggregator — Fantastic Furniture", "Week of 5 May 2026", "bhg.com.au (aggregator)", "Medium", "CA-9"),
        ("ASX — Super Retail Group (SUL)", "6 May 2026", "asx.com.au/asx/share-price-research/company/SUL", "High (ignored — no overlap)", "CA-10, CMP-1"),
        ("Investing.com — Nick Scali share price and analyst data", "May 2026", "investing.com", "High", "SO-8"),
        ("Market aggregators — EV/EBITDA multiples", "May 2026", "Multiple (unverified; see review flag)", "Medium", "CMP-2 through CMP-5"),
        ("Stock Titan — Tempur Sealy / Fullpower-AI partnership", "May 2026", "stocktitan.net", "Medium (AU applicability unconfirmed)", "CA (noted)"),
        ("Drewry World Container Index (WCI)", "30 Apr 2026", "drewry.co.uk/supply-chain-advisors/supply-chain-expertise/world-container-index", "High", "IG-1, IG-6"),
        ("EpicSourcing — ocean freight surcharge data", "May 2026", "epicsourcing.com.au", "High", "IG-6"),
        ("Australia Post — fuel surcharge update", "23 April 2026", "auspost.com.au", "Very High", "IG-6"),
        ("ABS Building Approvals / Dwelling Commencements", "Q3 2025 data", "abs.gov.au/statistics/industry/building-and-construction", "High", "IG-9"),
        ("Motley Fool AU — Harvey Norman analysis", "May 2026", "fool.com.au", "Medium-High (commentary, not primary)", "CA-3"),
    ]

    table_src = doc.add_table(rows=1, cols=5)
    table_src.style = "Table Grid"
    add_table_header_row(table_src, ["Source Name", "Date", "URL / Reference", "Confidence", "Used In"])
    for row_data in sources:
        row = table_src.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i == 2:
                run.font.color.rgb = MID_BLUE
                run.italic = True

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Approval record
    # ------------------------------------------------------------------
    styled_heading(doc, "5. Human Review and Approval Record", level=1)

    body_para(doc, "This research cycle has been reviewed and approved for document generation by the human reviewer.", bold=True)

    approval_table = doc.add_table(rows=5, cols=2)
    approval_table.style = "Table Grid"
    approval_data = [
        ("Approval status", "APPROVED — documents generated post approval"),
        ("Approval date", "8 May 2026"),
        ("Approving party", "Human reviewer (Hypnos Group / LK Group)"),
        ("Documents generated", "lkg-furniture-gm-weekly-digest.docx; lkg-furniture-board-weekly-digest.docx; lkg-furniture-internal-source-log.docx"),
        ("Classification applied", "GM / Board / Both / Ignored — as per classification summary above"),
    ]
    for i, (label, val) in enumerate(approval_data):
        row = approval_table.rows[i]
        r1 = row.cells[0].paragraphs[0].add_run(label)
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = row.cells[1].paragraphs[0].add_run(val)
        r2.font.size = Pt(9)

    doc.add_paragraph()
    disclaimer_footer(doc)

    out_path = os.path.join(OUTPUT_DIR, "lkg-furniture-internal-source-log.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}  ({os.path.getsize(out_path):,} bytes)")
    return out_path


# ===========================================================================
# MAIN
# ===========================================================================
if __name__ == "__main__":
    print("Generating LKG Furniture weekly digest documents...")
    print()
    gm_path    = build_gm_digest()
    board_path = build_board_digest()
    log_path   = build_source_log()
    print()
    print("All three documents generated successfully.")
    print(f"  GM Digest:      {gm_path}")
    print(f"  Board Digest:   {board_path}")
    print(f"  Source Log:     {log_path}")
